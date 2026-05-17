from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── 페이지 여백 설정 ──────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(3.0)
section.right_margin  = Cm(2.5)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)

# ── 스타일 헬퍼 ───────────────────────────────────────────────────────────────
def set_run_font(run, size=11, bold=False, color=None, name='맑은 고딕'):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    # 한글 폰트 설정
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), name)
    rPr.insert(0, rFonts)

def heading1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    set_run_font(run, size=16, bold=True, color=(31, 73, 125))
    # 아래 선
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), '1F497D')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def heading2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    set_run_font(run, size=13, bold=True, color=(0, 70, 127))
    return p

def heading3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    set_run_font(run, size=12, bold=True, color=(68, 114, 196))
    return p

def body(text, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Cm(0.8)
    run = p.add_run(text)
    set_run_font(run, size=11)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent  = Cm(0.8 + level * 0.6)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    set_run_font(run, size=11)
    return p

def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    # 헤더 행
    hdr_cells = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = hdr_cells[i].paragraphs[0].runs[0]
        set_run_font(run, size=10, bold=True, color=(255, 255, 255))
        tc = hdr_cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '2E75B6')
        tcPr.append(shd)
    # 데이터 행
    for ri, row_data in enumerate(rows):
        cells = t.rows[ri + 1].cells
        for ci, cell_text in enumerate(row_data):
            cells[ci].text = cell_text
            run = cells[ci].paragraphs[0].runs[0] if cells[ci].paragraphs[0].runs else cells[ci].paragraphs[0].add_run(cell_text)
            if cells[ci].paragraphs[0].runs:
                set_run_font(cells[ci].paragraphs[0].runs[0], size=10)
            fill = 'F2F2F2' if ri % 2 == 0 else 'FFFFFF'
            tc = cells[ci]._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), fill)
            tcPr.append(shd)
    if col_widths:
        for ri, row in enumerate(t.rows):
            for ci, cell in enumerate(row.cells):
                cell.width = Cm(col_widths[ci])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.8)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(30, 100, 30)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F0F0F0')
    pPr.append(shd)
    return p

def callout(text, color='FFF8DC'):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent   = Cm(0.6)
    p.paragraph_format.right_indent  = Cm(0.6)
    p.paragraph_format.space_before  = Pt(4)
    p.paragraph_format.space_after   = Pt(4)
    run = p.add_run(text)
    set_run_font(run, size=10.5)
    run.font.color.rgb = RGBColor(80, 60, 0)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    pPr.append(shd)
    return p

# ═══════════════════════════════════════════════════════════════════════════════
# 표지
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title_p.add_run('중간점검 보고서')
set_run_font(r, size=26, bold=True, color=(31, 73, 125))

doc.add_paragraph()
sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub_p.add_run('커스텀 악보 표기법 변환 및 연주 학습 앱')
set_run_font(r2, size=16, color=(68, 114, 196))

doc.add_paragraph()
date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = date_p.add_run('2026년 5월 14일')
set_run_font(r3, size=12, color=(127, 127, 127))

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 1. 팀별 솔루션 아이디어
# ═══════════════════════════════════════════════════════════════════════════════
heading1('1. 팀별 솔루션 아이디어')

# 1-1 키워드 카테고리
heading2('1-1. 키워드 카테고리 선택: 연결(Connection)')

callout(
    '"연결"이란 키워드는 음악을 즐기고 싶으나 악보라는 장벽으로 인해 진입하지 못한 사람들과 '
    '음악 사이의 간극을 좁힌다는 이 프로젝트의 본질을 가장 잘 대변한다.',
    color='EBF3FB'
)

doc.add_paragraph()
body('▶ "연결"을 선택한 이유')

bullet(
    '악보 문해력(음악 리터러시)은 단시간에 습득되지 않는다. '
    '오선보를 읽으려면 음표 위치, 박자, 조표 등 다양한 기호 체계를 동시에 학습해야 하므로 '
    '초보자에게는 높은 진입 장벽이 된다.'
)
bullet(
    '스마트폰이 보편화된 현재, 피아노를 배우고 싶어도 악보를 읽지 못해 포기하는 사용자가 '
    '여전히 다수 존재한다. 이 프로젝트는 그 사용자들을 음악과 "연결"시키는 다리 역할을 목표로 한다.'
)
bullet(
    '"변화"나 "예측"보다 "연결"이 더 적합한 이유: 이 앱의 핵심 가치는 기존 음악 콘텐츠(악보)를 '
    '변형하여 새로운 사용자 그룹에게 전달하는 것이다. 즉, 기존 자산(악보)과 신규 사용자(입문자) '
    '사이의 연결이 핵심이다.'
)

# 1-2 개발 배경 및 목적
heading2('1-2. 개발 배경 및 목적')

heading3('개발 배경')
body(
    '피아노 학습 시장에서 악보를 읽지 못하는 사용자는 대개 두 가지 경로를 택한다. '
    '첫째, 유튜브 등 동영상 미디어를 통해 보고 따라 치는 방식이다. 이는 학습 효율이 낮고 '
    '곡의 구조를 이해하기 어렵다. 둘째, 전문 선생님의 레슨을 통해 오선보를 처음부터 배우는 방식이다. '
    '이는 시간과 비용이 크게 소요된다.'
)
body(
    '본 프로젝트는 세 번째 경로를 제시한다. 기존 오선보를 촬영하거나 업로드하면 OMR(Optical Music '
    'Recognition) 기술이 음표 정보를 자동으로 인식하고, 직관적인 커스텀 표기법으로 재출력한다. '
    '사용자는 별도의 악보 교육 없이도 커스텀 악보를 읽고 피아노를 연주할 수 있다.'
)

heading3('개발 목적')
add_table(
    ['항목', '내용'],
    [
        ['핵심 문제', '일반 오선보는 악보 문해력 없이 읽기 어려우며 피아노 입문의 장벽으로 작용'],
        ['목표 사용자', '악보를 처음 접하거나 읽기 어려운 사람 (피아노 입문자 중심)'],
        ['접근 방식', 'OMR 기술로 기존 악보를 자동 인식 후 커스텀 표기법으로 변환 출력'],
        ['플랫폼', 'Android / iOS 모바일 앱 + 웹 온라인 데모'],
        ['라이선스', '학습 파이프라인 전체 상업 라이선스 확인 완료'],
    ],
    col_widths=[4, 11.5]
)

heading3('필요성')
bullet('악보 리터러시 없이 즉시 연주 시작 가능 → 음악 입문 장벽 제거')
bullet('기존 악보 자산(수백만 곡)을 변환 없이 활용 → 콘텐츠 확장성')
bullet('실시간 연주 피드백(정답/오답 표시) → 자기주도 학습 가능')
bullet('커스텀 악보 공유 커뮤니티 → 사용자 간 연결 생태계 형성')
bullet('향후 MIDI·마이크 연동 → 실제 피아노와의 연결로 확장')

# 1-3 솔루션 내용
heading2('1-3. 솔루션 내용')

heading3('(A) 온라인 체험 웹페이지')
body(
    '현재 운영 중인 웹 데모(web-demo/)는 별도 앱 설치 없이 브라우저에서 즉시 커스텀 악보와 '
    '피아노 연주를 체험할 수 있는 온라인 프로토타입이다. 순수 JavaScript(ES Modules)와 Web Audio API로 구현되어 있으며 서버 없이 정적 파일만으로 동작한다.'
)

add_table(
    ['구성 요소', '기능'],
    [
        ['튜토리얼', '커스텀 악보 3가지 규칙을 대화형으로 설명 (존 구조, 음가, 색상)'],
        ['악보 선택', '5종 샘플 악보(반짝반짝·학교종·기쁨의 송가 등) 썸네일 그리드'],
        ['커스텀 악보 렌더러', '선택한 악보를 Canvas에 커스텀 표기법으로 실시간 출력'],
        ['자동 재생', 'BPM 조절 가능한 메트로놈 기반 순차 재생, 음표 강조 애니메이션'],
        ['가상 피아노', '88건반 스크롤 UI, 다음 음 금색 강조, 오답 건반 빨간 지속 깜빡임'],
        ['참조 화살표', '높은음자리/낮은음자리 기준 건반 위치 화살표 표시'],
    ],
    col_widths=[4, 11.5]
)

heading3('(B) Flutter 모바일/웹 앱')
body(
    'Flutter 프레임워크를 사용하여 Android, iOS, 웹 브라우저를 단일 코드베이스로 지원한다. '
    '현재 브라우저 실행 가능한 UI 데모가 완성되어 있으며, 추후 OMR 엔진과 통합될 예정이다.'
)

add_table(
    ['탭', '화면', '기능'],
    [
        ['튜토리얼', '3규칙 PageView 위저드', '존/음가/색상 규칙 설명 + 음표 버튼 체험'],
        ['악보', '샘플 악보 그리드', '5종 선택 → CustomPainter로 커스텀 악보 출력'],
        ['연주', '악보+피아노 연동 화면', '순서대로 음 누르기, 진행률 바, 정답/오답 시각 피드백'],
    ],
    col_widths=[2, 4.5, 9]
)

heading3('(C) 커스텀 악보 표기법 (DeepScore)')
body(
    '오선보를 모르는 사용자도 즉시 읽을 수 있도록 설계된 독자적인 악보 표기 체계이다. '
    '총 3가지 핵심 규칙으로 구성된다.'
)

body('① 존(Zone) 구조 — 음높이 표현')
callout(
    '보표를 수직으로 3등분하여 위·중·아래 존(Zone)이 각각 다른 옥타브를 나타낸다.\n'
    '높은음자리(오른손) 3존 + 낮은음자리(왼손) 3존 = 총 6옥타브 커버.\n'
    '위 존이 높은 옥타브로, 피아노 건반의 배치 방향과 일치하여 직관적이다.',
    color='F0F8FF'
)
code_block(
    '높은음자리(오른손)\n'
    '┌────────────────────┐  ← 위 존  (C5~B5)\n'
    '│                    │\n'
    '├ ─ ─ ─ ─ ─ ─ ─ ─ ─┤  ← 중 존  (C4~B4)\n'
    '│                    │\n'
    '├ ─ ─ ─ ─ ─ ─ ─ ─ ─┤  ← 아래 존 (C3~B3)\n'
    '└────────────────────┘'
)

body('② 셀 너비 — 음가(박자) 표현')
callout(
    '한 마디를 박자수만큼 균등 셀로 분할하고, 음표는 자신의 음가에 비례하는 셀 수를 차지한다.\n'
    '예) 4/4박에서 2분음표 C + 4분음표 D + 4분음표 E:',
    color='F0FFF0'
)
code_block(
    '┌────────┬────┬────┐\n'
    '│   C    │ D  │ E  │\n'
    '└────────┴────┴────┘\n'
    '  ←2셀→   1셀  1셀'
)

body('③ 색상 테두리 — 박자 순서 표현')
callout(
    '음표 순서에 따라 주황(#FF8C00) → 연두(#7CFC00) → 하늘(#00BFFF) 테두리가 순환 적용된다.\n'
    '흰 건반은 영문 대문자(C D E F G A B), 검은 건반은 숫자(1~5)로 표기한다.',
    color='FFFAF0'
)

add_table(
    ['건반 종류', '표기', '예시'],
    [
        ['흰 건반', '영문 대문자', 'C, D, E, F, G, A, B'],
        ['검은 건반', '숫자 1~5', '1=C#/Db, 2=D#/Eb, 3=F#/Gb, 4=G#/Ab, 5=A#/Bb'],
    ],
    col_widths=[3.5, 4, 8]
)

heading3('(D) OMR 기반 악보 인식 모델 — 이론 설명')
body(
    'OMR(Optical Music Recognition)은 악보 이미지에서 음표·박자·조표 등의 음악 기호를 '
    '자동으로 인식하고 구조화된 데이터로 변환하는 기술이다. 본 프로젝트의 OMR 파이프라인은 '
    '세 개의 딥러닝 모델로 구성된 Encoder-Decoder 구조를 채택한다.'
)

body('① 세그멘테이션 모델 (TFLite)')
bullet(
    'SegNet 계열 아키텍처를 INT8 양자화하여 TFLite 포맷으로 변환 (segnet_308_int8.tflite).'
)
bullet(
    '입력: 악보 이미지 → 출력: 음표·오선 등 기호별 픽셀 세그멘테이션 맵.'
)
bullet(
    '전처리 단계: 자동 크롭 → 크기 조정 → 색상 보정 → 노이즈 제거 → 오선 검출 및 보정.'
)

body('② 인코더 모델 (TFLite)')
bullet(
    '세그멘테이션 결과를 입력으로 받아 악보의 시각적 특징을 고차원 벡터 시퀀스로 압축한다 (encoder_331_int8.tflite).'
)
bullet(
    'CNN 기반 특징 추출 후 Transformer 인코더를 통해 음표 간 문맥 관계를 학습한다.'
)

body('③ 디코더 모델 (ONNX Runtime)')
bullet(
    'Transformer 디코더 구조로 인코더 출력을 받아 음표 토큰 시퀀스를 순차 생성한다 (decoder_331_int8.onnx).'
)
bullet(
    '출력 토큰은 DeepScore 형식으로 정의된 어휘(Vocabulary) 1,012개 중 하나를 선택한다.'
)
bullet(
    '예) clef-G → key-C → time-4/4 → note-C4-1/4 → note-E4-1/4 → ... → barline-final'
)

body('④ 학습 데이터 및 어휘 체계')
bullet(
    '학습 데이터는 music21 라이브러리와 MuseScore(.mscz) 파일을 활용해 합성 생성한다. '
    '실제 악보를 구매하거나 저작권 침해 없이 무제한 데이터 생성이 가능하다.'
)
bullet(
    f'어휘(Vocabulary): 총 1,012 토큰 — 음표(850), 화음 추가음(85), 쉼표(10), '
    '세로줄(5), 셈여림(11), 아티큘레이션(5), 셋잇단음표·옥타브 기호 등 포함.'
)
bullet(
    '학습 단계: Round 1(기본 음표 3,000장) → Round 2(전체 기호 확장) → '
    'Round 3(그랜드 스태프) → Round 4(실사 사진 fine-tuning).'
)

heading3('(E) 학습된 엔진의 앱·웹 적용 방식')
body(
    '학습 완료된 모델을 실제 사용자 환경(앱·웹)에 배포하는 방식은 플랫폼별로 다르게 설계되어 있다.'
)

body('① Android 앱 (현재 구현 방식)')
bullet('학습된 모델(.tflite, .onnx)을 앱 assets 폴더에 번들링하여 배포.')
bullet(
    'Flutter → Kotlin(MethodChannel) → JNI → C++ 추론 엔진 순서로 호출. '
    'C++ 엔진이 OpenCV로 이미지를 전처리하고 TFLite/ONNX Runtime으로 추론을 수행한다.'
)
bullet(
    '추론 결과(DeepScore 토큰 시퀀스) → Python 렌더러 로직을 Dart로 포팅한 '
    'CustomPainter → 커스텀 악보 이미지로 화면에 출력.'
)

code_block(
    'Flutter UI (Dart)\n'
    '  ↓ MethodChannel\n'
    'Kotlin (OmrPipeline.kt)\n'
    '  ↓ JNI\n'
    'C++ 추론 엔진 (flutter_jni.cpp)\n'
    '  ↓\n'
    'TFLite (세그멘테이션·인코더) + ONNX Runtime (디코더)\n'
    '  ↓\n'
    'DeepScore 토큰 시퀀스 → Dart CustomPainter → 커스텀 악보 출력'
)

body('② iOS 앱 (예정 방식 — Phase 5 이후)')
bullet(
    'Objective-C++ 브리지(flutter_ios.mm)를 통해 동일한 C++ 엔진을 NSBundle 에셋 로더로 연결. '
    'CocoaPods에 OpenCV·TFLite·ONNX Runtime 추가. Android와 동일한 Dart 레이어 재사용.'
)

body('③ 웹 서비스 (FastAPI 방식)')
bullet(
    '브라우저에서 C++ WASM 빌드는 모델 크기·로딩 지연 문제로 비현실적. '
    '대신 FastAPI 서버가 OMR 추론 API 역할을 담당한다.'
)
bullet(
    '사용자 흐름: 브라우저에서 악보 이미지 업로드 → FastAPI 서버에서 C++ 엔진 추론 → '
    'DeepScore JSON 반환 → 브라우저 JavaScript 렌더러가 커스텀 악보 출력.'
)
bullet(
    '학습 전 단계(현재)에서는 MusicXML/mscz 파일을 직접 입력으로 받아 '
    '동일한 렌더러 파이프라인을 거쳐 커스텀 악보를 출력한다.'
)

code_block(
    '브라우저 (JavaScript)\n'
    '  ↓ HTTP POST /omr  (이미지 업로드)\n'
    'FastAPI 서버 (Python)\n'
    '  ↓\n'
    'C++ 추론 엔진 (subprocess 또는 ctypes)\n'
    '  ↓ DeepScore JSON 반환\n'
    '브라우저 Canvas 렌더러 → 커스텀 악보 표시'
)

body('④ 모델 업데이트 전략')
bullet(
    'Round별 학습 완료 시 모델 파일(.tflite, .onnx)만 교체하면 앱·웹 코드 변경 없이 '
    '정확도가 향상된다. 엔진-UI 레이어 분리 설계 덕분에 지속적인 모델 개선이 가능하다.'
)
bullet(
    '향후 Dart FFI로 JNI를 대체할 경우 Android/iOS 통합 코드베이스 유지가 가능해진다.'
)

# ═══════════════════════════════════════════════════════════════════════════════
# 2. 향후 개발 방향성 및 예상 결과물
# ═══════════════════════════════════════════════════════════════════════════════
heading1('2. 향후 개발 방향성 및 예상 결과물 형태')

heading2('2-1. 현재 구현 상태')
add_table(
    ['구성 요소', '상태', '비고'],
    [
        ['웹 데모 프로토타입', '✅ 완료', '튜토리얼·피아노·악보 렌더러·연주 체험'],
        ['Python 커스텀 악보 렌더러', '✅ 완료', 'render_notation.py'],
        ['C++ 추론 엔진', '✅ 완료', '학습 완료 후 모델 파일 교체 필요'],
        ['학습 데이터 생성 파이프라인', '✅ 완료', 'Vocabulary 1,012 토큰, Round 2 기호 포함'],
        ['Flutter 앱 UI 데모 (브라우저)', '✅ 완료', '튜토리얼·악보 선택·연주 체험 3탭'],
        ['Android Kotlin JNI 브리지', '✅ 완료', '—'],
        ['OMR 모델 학습 (Round 1~4)', '🔄 예정', 'RTX 3080, 기본 음표 3,000장부터'],
        ['iOS OMR 브리지', '❌ 예정', 'Objective-C++ 방식, Phase 5 이후'],
        ['실시간 카메라 입력', '❌ 예정', '—'],
        ['악보 공유 백엔드', '❌ 예정', 'Firebase / Supabase 검토 중'],
        ['실시간 연주 감지', '❌ 예정', 'Web MIDI API 우선'],
    ],
    col_widths=[5.5, 2, 8]
)

heading2('2-2. 단계별 개발 로드맵')
add_table(
    ['단계', '내용', '상태'],
    [
        ['Phase 0~1', '커스텀 표기법 확정 + 데이터 생성 파이프라인', '✅ 완료'],
        ['Phase 2', 'OMR 모델 학습 (Round 1~4, RTX 3080)', '🔄 진행 예정'],
        ['Phase 3', '모델 변환·양자화 (PyTorch → TFLite INT8)', '예정'],
        ['Phase 4', 'C++ 추론 엔진 구현', '✅ 완료'],
        ['Phase 5-A', 'Flutter 앱 UI 데모 (브라우저/데스크탑)', '✅ 완료'],
        ['Phase 5', 'Flutter 통합 (Dart FFI, 카메라, Riverpod)', '예정'],
        ['Phase 6', 'Flutter CustomPainter 렌더러 포팅', '예정'],
        ['Phase 7', '멀티플랫폼 완성 (iOS + 웹 FastAPI)', '예정'],
        ['Phase 8', '웹 고도화 (악보 업로드 → 실시간 변환 + 연주)', '예정'],
        ['Phase 9', '공유 커뮤니티 + MIDI/마이크 연주 감지', '예정'],
    ],
    col_widths=[2.5, 9, 2]
)

heading2('2-3. 예상 결과물 형태')

heading3('단기 (학습 완료 전 데모)')
bullet('브라우저에서 즉시 실행 가능한 Flutter 웹 앱 (flutter run -d chrome)')
bullet('샘플 악보 5종 선택 → 커스텀 악보 출력 → 피아노 건반 연주 체험')
bullet('웹 데모 온라인 배포 (GitHub Pages 또는 Netlify 예정)')

heading3('중기 (OMR 모델 학습 완료 후)')
bullet('스마트폰 카메라로 악보 촬영 → 수초 내 커스텀 악보 변환 (Android 앱)')
bullet('FastAPI 서버 연동 웹 서비스: 악보 이미지 업로드 → 변환 결과 출력')
bullet('iOS 앱 배포 (Objective-C++ 브리지 구현 후)')

heading3('장기 (최종 목표)')
bullet('MIDI 키보드 연결 → 실시간 연주 정확도 피드백')
bullet('커스텀 악보 공유 커뮤니티 (Firebase / Supabase 백엔드)')
bullet('연주 점수 표시, 악보 난이도 자동 태깅, 맞춤형 학습 추천')

heading2('2-4. 활용 기술 스택')
add_table(
    ['영역', '기술'],
    [
        ['모바일 앱 UI', 'Flutter (Dart 3), CustomPainter, Riverpod'],
        ['OMR 모델 학습', 'PyTorch, music21, MuseScore(.mscz)'],
        ['추론 엔진', 'C++20, OpenCV 4.9, TFLite 2.16, ONNX Runtime 1.18'],
        ['Android 브리지', 'Kotlin JNI (현재) → Dart FFI (예정)'],
        ['iOS 브리지', 'Objective-C++ + MethodChannel'],
        ['웹 프론트엔드', 'Vanilla JS + Web Audio API (데모) / Flutter Web (앱)'],
        ['웹 백엔드', 'FastAPI + Python 렌더러 (OMR API)'],
        ['학습 데이터', 'music21, MuseScore, 실사 악보 사진'],
        ['공유 백엔드', 'Firebase / Supabase'],
        ['연주 입력', 'Web MIDI API (1순위) / Web Audio API 단음 감지 (2순위)'],
    ],
    col_widths=[4, 11.5]
)

# ── 저장 ──────────────────────────────────────────────────────────────────────
out = r'C:\Users\kyutae\AndroidStudioProjects\musicscore_flutter\중간점검_보고서.docx'
doc.save(out)
print(f'saved: {out}')
